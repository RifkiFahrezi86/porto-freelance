from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document(r'd:\SUPERVISOR\SPORT\Template-buku-pjkr.docx')

# ======== 1. FIX BASE FONTS TO TIMES NEW ROMAN ========
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)

isi_style = doc.styles['isi']
isi_style.font.name = 'Times New Roman'
isi_style.font.size = Pt(12)

doc.styles['Heading 1'].font.name = 'Times New Roman'
doc.styles['Heading 2'].font.name = 'Times New Roman'
doc.styles['Heading 2'].font.size = Pt(12)

try:
    doc.styles['toc 1'].font.name = 'Times New Roman'
    doc.styles['toc 1'].font.size = Pt(12)
    doc.styles['toc 2'].font.name = 'Times New Roman'
    doc.styles['toc 2'].font.size = Pt(12)
except:
    pass

# ======== 2. UPDATE COVER PAGE ========
p0 = doc.paragraphs[0]
for run in p0.runs:
    run.text = ""
p0.runs[0].text = "MENGANALISIS KETERAMPILAN GERAK\nPERMAINAN BOLA BESAR\n(SEPAK BOLA)"
for run in p0.runs:
    run.font.name = 'Times New Roman'

p5 = doc.paragraphs[5]
for run in p5.runs:
    run.text = ""
p5.runs[0].text = "Kelompok 1"
for run in p5.runs:
    run.font.name = 'Times New Roman'

p6 = doc.paragraphs[6]
for run in p6.runs:
    run.text = ""
p6.runs[0].text = "Nur Arslan\nSyahrul\nFadhel Akbar\nDwi Cahya Mentari"
for run in p6.runs:
    run.font.name = 'Times New Roman'

# ======== 3. REMOVE CONTENT FROM PARAGRAPH 23 ONWARD ========
body = doc.element.body
elements_to_remove = []
para_count = 0
for child in list(body):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        if para_count >= 23:
            elements_to_remove.append(child)
        para_count += 1
    elif tag == 'tbl' and para_count >= 23:
        elements_to_remove.append(child)

for el in elements_to_remove:
    body.remove(el)

# ======== HELPER FUNCTIONS ========
def add_heading1(text):
    p = doc.add_paragraph(style='Heading 1')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    return p

def add_heading2(text):
    """Format: 1.1.   JUDUL SUB BAB (Heading 2 style, bold)"""
    p = doc.add_paragraph(style='Heading 2')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_isi(text):
    p = doc.add_paragraph(style='isi')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_sub_point(text):
    """Sub-point within isi (bold, no first-line indent)"""
    p = doc.add_paragraph(style='isi')
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_page_break():
    doc.add_page_break()

def add_section_break():
    return doc.add_section()

def create_footer_with_page_number(section):
    # Matikan 'Different First Page' agar semua halaman punya footer
    sectPr = section._sectPr
    titlePg = sectPr.find(qn('w:titlePg'))
    if titlePg is not None:
        sectPr.remove(titlePg)

    footer = section.footer
    footer.is_linked_to_previous = False
    # Bersihkan footer lama
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # PAGE field
    run2 = p.add_run()
    fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
    run2._element.append(fldChar1)

    run3 = p.add_run()
    instrText = parse_xml(r'<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(nsdecls('w')))
    run3._element.append(instrText)

    run4 = p.add_run()
    fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls('w')))
    run4._element.append(fldChar2)

    run5 = p.add_run("1")
    run5.font.name = 'Times New Roman'
    run5.font.size = Pt(10)

    run6 = p.add_run()
    fldChar3 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
    run6._element.append(fldChar3)

    # Juga bersihkan first page footer jika ada
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    for p_f in first_footer.paragraphs:
        p_f.clear()
    p_f2 = first_footer.paragraphs[0]
    p_f2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p_f2.add_run()
    r2._element.append(parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w'))))
    r3 = p_f2.add_run()
    r3._element.append(parse_xml(r'<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(nsdecls('w'))))
    r4 = p_f2.add_run()
    r4._element.append(parse_xml(r'<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls('w'))))
    r5 = p_f2.add_run("1")
    r5.font.name = 'Times New Roman'
    r5.font.size = Pt(10)
    r6 = p_f2.add_run()
    r6._element.append(parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w'))))

def set_page_number_format(section, fmt='decimal', start=None):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = parse_xml('<w:pgNumType {}/>'.format(nsdecls('w')))
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))

# ======== SECTION 0: SAMPUL (tanpa nomor halaman) ========
sec0 = doc.sections[0]
# Matikan titlePg supaya tidak inherit aneh
titlePg0 = sec0._sectPr.find(qn('w:titlePg'))
if titlePg0 is not None:
    sec0._sectPr.remove(titlePg0)
sec0.footer.is_linked_to_previous = False
for p_f in sec0.footer.paragraphs:
    p_f.clear()
sec0.first_page_footer.is_linked_to_previous = False
for p_f in sec0.first_page_footer.paragraphs:
    p_f.clear()

# ======== SECTION BREAK: DAFTAR ISI (angka Romawi) ========
sec_daftar = add_section_break()
set_page_number_format(sec_daftar, fmt='lowerRoman', start=1)
create_footer_with_page_number(sec_daftar)

# ======== DAFTAR ISI ========
add_heading1("DAFTAR ISI")

# TOC field
p_toc = doc.add_paragraph()
run_toc = p_toc.add_run()
run_toc._element.append(parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w'))))
run_toc2 = p_toc.add_run()
run_toc2._element.append(parse_xml(r'<w:instrText {} xml:space="preserve"> TOC \o "1-2" \h \z \u </w:instrText>'.format(nsdecls('w'))))
run_toc3 = p_toc.add_run()
run_toc3._element.append(parse_xml(r'<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls('w'))))

toc_entries = [
    ("toc 1", "DAFTAR ISI\tii"),
    ("toc 1", "BAB 1 MENGANALISIS KETERAMPILAN GERAK PERMAINAN BOLA BESAR (SEPAK BOLA)\t1"),
    ("toc 2", "1.1.\tPengertian dan Konsep Dasar Sepak Bola\t1"),
    ("toc 2", "1.2.\tAnalisis Gerak Menendang dengan Kaki Bagian Dalam\t2"),
    ("toc 2", "1.3.\tAnalisis Gerak Menendang dengan Kaki Bagian Luar\t4"),
    ("toc 2", "1.4.\tAnalisis Gerak Menendang dengan Punggung Kaki\t6"),
    ("toc 2", "1.5.\tPerbandingan Ketiga Teknik Tendangan\t8"),
    ("toc 2", "1.6.\tRingkasan\t9"),
    ("toc 2", "1.7.\tPenilaian\t10"),
    ("toc 1", "DAFTAR PUSTAKA\t11"),
]
for style_name, entry_text in toc_entries:
    p_entry = doc.add_paragraph(style=style_name)
    run_entry = p_entry.add_run(entry_text)
    run_entry.font.name = 'Times New Roman'
    run_entry.font.size = Pt(12)

run_tocend = doc.add_paragraph().add_run()
run_tocend._element.append(parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w'))))

# ======== SECTION BREAK: BAB 1 (Arabic numbers from 1) ========
sec_bab = add_section_break()
set_page_number_format(sec_bab, fmt='decimal', start=1)
create_footer_with_page_number(sec_bab)

# ================================================================
# BAB 1
# ================================================================

add_heading1("BAB 1\nMENGANALISIS KETERAMPILAN GERAK PERMAINAN\nBOLA BESAR (SEPAK BOLA)")

add_isi(
    "Di antara sekian banyak materi PJOK yang diajarkan di SMA, MA, SMK, maupun MAK, sepak bola sudah pasti jadi salah satu yang paling ditunggu-tunggu oleh siswa. Hampir semua orang pernah main sepak bola, minimal di halaman rumah atau lapangan kampung. Tapi yang menarik, banyak dari kita yang sudah terbiasa menendang bola sejak kecil tanpa benar-benar tahu teknik mana yang sedang kita pakai. Padahal dalam sepak bola, cara kita menendang bola itu menentukan segalanya — mulai dari arah operan, kecepatan bola, sampai akurasi tembakan ke gawang. Bab ini akan fokus menganalisis teknik-teknik tendangan dalam sepak bola, khususnya tendangan menggunakan kaki bagian dalam, kaki bagian luar, dan punggung kaki (Mahendra, 2020)."
)

# ================================================================
# 1.1. PENGERTIAN DAN KONSEP DASAR
# ================================================================

add_heading2("1.1.\tPengertian dan Konsep Dasar Sepak Bola")

add_isi(
    "Sepak bola adalah olahraga beregu yang mempertemukan dua tim, masing-masing sebelas pemain di atas lapangan. Tujuannya sederhana: masukkan bola ke gawang lawan sebanyak-banyaknya. Pemain tidak boleh memakai tangan kecuali kiper di area kotak penaltinya sendiri. Luxbacher (2011) menjelaskan bahwa sepak bola menuntut perpaduan antara fisik, teknik, dan pemahaman taktis dari setiap pemainnya."
)

add_sub_point("a. Menendang sebagai Keterampilan Utama")
add_isi(
    "Kalau ditanya teknik apa yang paling fundamental di sepak bola, jawabannya pasti menendang bola. Hampir semua aspek permainan melibatkan tendangan — mulai dari mengoper ke rekan setim, menembak ke gawang, sampai melakukan tendangan bebas dan corner kick. Mielke (2007) bahkan menyebut bahwa menendang bola adalah keterampilan pertama yang harus dikuasai oleh siapa pun yang ingin bermain sepak bola. Tanpa kemampuan menendang yang baik, sehebat apa pun fisik dan stamina seorang pemain, dia akan kesulitan berkontribusi di lapangan."
)

add_sub_point("b. Bagian Kaki yang Digunakan untuk Menendang")
add_isi(
    "Secara umum, ada tiga bagian kaki yang paling sering dipakai untuk menendang bola dalam sepak bola. Pertama, kaki bagian dalam atau inside foot, yang terletak di sisi medial kaki — area tulang di bawah mata kaki bagian dalam. Kedua, kaki bagian luar atau outside foot, yang terletak di sisi lateral kaki — area tulang di bawah mata kaki bagian luar. Ketiga, punggung kaki atau instep, yaitu bagian atas kaki di mana tali sepatu berada. Masing-masing bagian ini menghasilkan karakteristik tendangan yang berbeda dari segi kekuatan, akurasi, dan putaran bola. Dalam sub-bab berikutnya, ketiga teknik tendangan ini akan dianalisis satu per satu secara mendalam (Sucipto, 2015)."
)

# ================================================================
# 1.2. TENDANGAN KAKI BAGIAN DALAM
# ================================================================

add_heading2("1.2.\tAnalisis Gerak Menendang dengan Kaki Bagian Dalam")

add_isi(
    "Tendangan menggunakan kaki bagian dalam ini bisa dibilang teknik yang paling sering dipakai di semua level permainan, dari pemula sampai profesional. Kenapa? Karena area perkenaan kaki bagian dalam itu relatif lebar dan datar, sehingga kontrol terhadap arah bola jadi lebih mudah. Hampir semua operan pendek dan menengah di sepak bola menggunakan teknik ini. Bahkan beberapa pemain top dunia juga mengandalkan kaki bagian dalam untuk tendangan bebas karena bisa menghasilkan putaran bola yang indah."
)

add_sub_point("a. Sikap Awal dan Posisi Tubuh")
add_isi(
    "Sebelum menendang, ada beberapa hal yang perlu diperhatikan soal posisi tubuh. Kaki tumpuan ditempatkan di samping bola, jaraknya kira-kira satu kepalan tangan, dengan ujung kaki mengarah ke target. Lutut kaki tumpuan sedikit ditekuk supaya badan lebih stabil dan tidak goyah saat kaki ayun menendang. Badan agak condong ke depan dan pandangan mata fokus ke bola, bukan ke target — karena kalau mata sudah fokus ke target sebelum menendang, biasanya justru tendangannya meleset. Kedua tangan direntangkan ke samping buat menjaga keseimbangan (Sucipto, 2015)."
)

add_sub_point("b. Fase Ayunan dan Perkenaan Bola")
add_isi(
    "Kaki yang menendang diayunkan ke belakang dulu sebagai ancang-ancang, lalu diayun ke depan mengenai bola. Nah, di sinilah bagian pentingnya: perkenaan bola harus tepat di tengah-tengah bola menggunakan area datar kaki bagian dalam, yaitu tulang yang ada di bawah mata kaki bagian dalam. Kalau perkenaannya terlalu tinggi, bola akan menyusur tanah. Kalau terlalu rendah, bola malah naik ke atas. Pergelangan kaki harus dikunci dan tidak boleh kendor supaya tenaga dari ayunan kaki tersalurkan penuh ke bola. Satu kesalahan yang sering terjadi pada siswa pemula adalah pergelangan kaki yang terlalu lemas, akibatnya bola tidak bertenaga dan arahnya tidak terkontrol (Mielke, 2007)."
)

add_sub_point("c. Gerakan Lanjutan (Follow-through)")
add_isi(
    "Setelah kaki mengenai bola, gerakan kaki tidak boleh langsung berhenti. Kaki tetap mengayun ke depan mengikuti arah tendangan — ini yang disebut follow-through. Fungsinya penting sekali: follow-through memastikan tenaga yang dihasilkan optimal dan arah bola sesuai dengan yang diinginkan. Tanpa follow-through yang baik, tendangan sering kali kurang bertenaga atau arahnya berbelok dari target. Badan juga ikut bergerak sedikit ke depan mengikuti momentum tendangan."
)

add_sub_point("d. Kapan Menggunakan Tendangan Kaki Bagian Dalam")
add_isi(
    "Teknik ini paling cocok dipakai untuk situasi-situasi berikut: operan pendek dan menengah ke rekan setim karena akurasinya tinggi, tendangan penalti karena arah bola bisa dikontrol dengan presisi, serta tendangan bebas jarak dekat hingga menengah. Di level profesional, banyak pemain seperti Andrea Pirlo atau Lionel Messi yang terkenal dengan tendangan bebas pakai kaki bagian dalam yang menghasilkan curl atau lengkungan bola yang luar biasa. Kelemahannya, tendangan ini biasanya tidak menghasilkan kecepatan sebesar tendangan pakai punggung kaki, jadi kurang cocok buat tembakan jarak jauh yang butuh power besar (Luxbacher, 2011)."
)

# ================================================================
# 1.3. TENDANGAN KAKI BAGIAN LUAR
# ================================================================

add_heading2("1.3.\tAnalisis Gerak Menendang dengan Kaki Bagian Luar")

add_isi(
    "Kalau tendangan kaki bagian dalam itu ibarat teknik yang polos dan bisa diandalkan, maka tendangan kaki bagian luar ini lebih ke arah teknik yang tricky dan tidak terduga. Banyak pemain top dunia yang punya signature move pakai kaki bagian luar — sebut saja Roberto Carlos dengan tendangan bebas legendaris-nya, atau Quaresma dengan trivela yang bikin kagum penonton. Dalam pembelajaran di sekolah, teknik ini memang jarang diajarkan secara mendalam, padahal sebenarnya sangat berguna di situasi-situasi tertentu."
)

add_sub_point("a. Sikap Awal dan Posisi Tubuh")
add_isi(
    "Posisi awal untuk tendangan kaki bagian luar sedikit berbeda dari kaki bagian dalam. Kaki tumpuan diletakkan agak di belakang dan sedikit ke samping bola, bukan tepat di samping seperti tendangan kaki bagian dalam. Badan agak dicondongkan ke arah kaki tumpuan, jadi sedikit miring menjauhi bola. Pandangan tetap ke bola dan kedua tangan dipakai buat menjaga keseimbangan. Posisi ini memang agak kurang natural dibanding tendangan kaki bagian dalam, makanya banyak pemula yang merasa canggung saat pertama kali mencobanya (Scheunemann, 2014)."
)

add_sub_point("b. Fase Ayunan dan Perkenaan Bola")
add_isi(
    "Kaki yang menendang diayunkan dari belakang ke depan, tapi bedanya — perkenaan dengan bola terjadi di bagian luar kaki, yaitu area tulang di bawah mata kaki bagian luar. Pergelangan kaki harus diputar sedikit ke dalam (inversi) supaya permukaan kaki bagian luar membentuk bidang datar yang cukup untuk menyentuh bola. Yang perlu diingat, titik perkenaan pada bola juga berbeda: kalau mau bola lurus, perkenaannya di tengah bola; kalau mau bola melengkung, perkenaannya agak ke sisi dalam bola sehingga menghasilkan putaran ke arah luar. Mielke (2007) mengingatkan bahwa tendangan ini butuh latihan yang cukup banyak karena area perkenaan kaki bagian luar lebih kecil dibanding kaki bagian dalam."
)

add_sub_point("c. Gerakan Lanjutan (Follow-through)")
add_isi(
    "Follow-through pada tendangan kaki bagian luar punya karakteristik yang khas. Setelah mengenai bola, kaki lanjut bergerak menyilang ke arah dalam tubuh, tidak lurus ke depan seperti tendangan kaki bagian dalam. Gerakan menyilang ini secara alami menghasilkan putaran atau spin pada bola yang membuat lintasannya melengkung. Besarnya lengkungan tergantung dari seberapa kuat pergelangan kaki diputar dan seberapa jauh kaki menyilang saat follow-through."
)

add_sub_point("d. Kapan Menggunakan Tendangan Kaki Bagian Luar")
add_isi(
    "Tendangan kaki bagian luar paling efektif di beberapa situasi. Pertama, untuk operan cepat ke samping tanpa harus mengubah posisi tubuh secara drastis — misalnya saat berlari ke depan tapi ingin mengoper ke pemain yang ada di samping. Kedua, untuk memberikan umpan melengkung yang sulit dibaca oleh pemain bertahan lawan. Ketiga, di beberapa situasi shooting di mana sudut tendangan terbatas dan pemain tidak punya waktu untuk mengatur posisi tubuh buat menendang pakai kaki bagian dalam. Kelemahannya, akurasi tendangan ini umumnya lebih rendah dan butuh latihan lebih banyak untuk bisa konsisten. Selain itu, kekuatan tendangan juga biasanya tidak sebesar teknik lainnya karena otot-otot yang terlibat berbeda (Sucipto, 2015)."
)

# ================================================================
# 1.4. TENDANGAN PUNGGUNG KAKI
# ================================================================

add_heading2("1.4.\tAnalisis Gerak Menendang dengan Punggung Kaki")

add_isi(
    "Kalau kita nonton pertandingan dan ada pemain yang melepaskan tembakan keras yang melesat ke pojok gawang, hampir pasti dia pakai punggung kaki. Teknik ini menghasilkan tendangan paling keras di antara ketiga teknik yang kita bahas karena memanfaatkan area kaki yang keras dan luas serta melibatkan otot-otot tungkai yang lebih besar. Tidak heran kalau sebagian besar gol dari tembakan jarak jauh menggunakan teknik tendangan punggung kaki."
)

add_sub_point("a. Sikap Awal dan Posisi Tubuh")
add_isi(
    "Ancang-ancang untuk tendangan punggung kaki biasanya lebih panjang dibanding dua teknik sebelumnya, sekitar dua sampai tiga langkah mundur dari bola. Kaki tumpuan ditempatkan di samping bola dengan jarak yang nyaman, ujung kaki mengarah lurus ke target. Badan agak condong ke depan — kalau terlalu tegak atau malah condong ke belakang, bola cenderung melayang tinggi tidak terkontrol. Pandangan fokus ke bola dan kepala agak menunduk. Ini detail kecil yang sering diabaikan: kalau kepala mendongak ke atas saat menendang, badan ikut terangkat dan bola hampir pasti melambung melewati gawang (Luxbacher, 2011)."
)

add_sub_point("b. Fase Ayunan dan Perkenaan Bola")
add_isi(
    "Ayunan kaki untuk tendangan punggung kaki dimulai dari pinggul, bukan dari lutut saja. Ini yang membedakannya dari dua teknik sebelumnya — karena melibatkan ayunan dari pinggul, tenaga yang dihasilkan jauh lebih besar. Kaki diayun dari belakang dengan lutut ditekuk, kemudian saat mendekati bola lutut diluruskan dengan cepat dan kuat. Perkenaan bola terjadi di punggung kaki, yaitu area atas kaki di mana tali sepatu berada. Pergelangan kaki harus dikunci kuat dalam posisi plantar fleksi, artinya ujung kaki menunjuk ke bawah. Jari-jari kaki juga ditekuk ke bawah. Bola yang ditendang dengan benar akan melesat lurus dan keras. Tapi kalau pergelangan kaki kendor, bola bisa meleset kemana-mana dan bahkan bisa menyebabkan cedera pada kaki penendang (Mielke, 2007)."
)

add_sub_point("c. Gerakan Lanjutan (Follow-through)")
add_isi(
    "Follow-through untuk tendangan punggung kaki sangat berpengaruh terhadap kecepatan akhir bola. Setelah mengenai bola, kaki tetap bergerak ke atas dan ke depan mengikuti jalur tendangan. Semakin penuh follow-through-nya, semakin besar transfer energi ke bola. Badan ikut terbawa ke depan, dan kadang kaki tumpuan bahkan sedikit terangkat dari tanah akibat momentum yang dihasilkan. Di sinilah pentingnya keseimbangan — pemain yang kurang kuat core muscle-nya sering kehilangan keseimbangan setelah melakukan tendangan punggung kaki yang keras."
)

add_sub_point("d. Kapan Menggunakan Tendangan Punggung Kaki")
add_isi(
    "Teknik ini paling ideal untuk: tembakan ke gawang dari jarak menengah sampai jauh, tendangan bebas langsung, operan jarak jauh (long pass) untuk memindahkan bola dari satu sisi lapangan ke sisi lainnya, dan tendangan voli atau half-volley. Kelebihannya jelas — power yang dihasilkan paling besar di antara ketiga teknik. Tapi kekurangannya, akurasi biasanya lebih rendah dibanding tendangan kaki bagian dalam, dan teknik ini juga lebih sulit dikuasai oleh pemula karena area perkenaan yang harus tepat. Sedikit saja meleset dari punggung kaki, bola bisa menyimpang jauh dari target (Sucipto, 2015)."
)

# ================================================================
# 1.5. PERBANDINGAN KETIGA TEKNIK
# ================================================================

add_heading2("1.5.\tPerbandingan Ketiga Teknik Tendangan")

add_isi(
    "Setelah membahas satu per satu, sekarang kita coba bandingkan ketiga teknik tendangan ini biar lebih jelas perbedaan dan kegunaannya masing-masing."
)

add_sub_point("a. Dari Segi Akurasi")
add_isi(
    "Kalau soal akurasi, tendangan kaki bagian dalam jelas unggul. Area perkenaannya yang datar dan lebar membuat pemain lebih mudah mengontrol arah bola. Tendangan punggung kaki ada di posisi kedua, karena meskipun area perkenaannya cukup luas, power yang besar kadang mengorbankan akurasi. Tendangan kaki bagian luar paling sulit dari segi akurasi karena area perkenaannya paling kecil dan posisi tubuhnya kurang natural."
)

add_sub_point("b. Dari Segi Kekuatan")
add_isi(
    "Sebaliknya, kalau bicara soal kekuatan tendangan, punggung kaki jadi juaranya. Ayunan yang melibatkan seluruh otot tungkai dari pinggul sampai ujung kaki menghasilkan power paling besar. Kaki bagian luar ada di posisi kedua karena bisa menghasilkan tendangan yang cukup keras dengan putaran tambahan. Kaki bagian dalam menghasilkan kekuatan paling rendah karena gerakan ayunannya lebih terbatas oleh posisi kaki yang terbuka ke samping."
)

add_sub_point("c. Dari Segi Putaran Bola")
add_isi(
    "Soal putaran atau spin bola, tendangan kaki bagian luar punya keunikan tersendiri karena menghasilkan putaran swerve yang melengkung ke arah luar — ini yang sering disebut teknik trivela. Kaki bagian dalam juga bisa menghasilkan putaran, tapi arahnya ke dalam (curl). Punggung kaki secara default menghasilkan tendangan dengan putaran minimal alias lurus, meskipun pemain terampil bisa memberikan sedikit sidespin dengan mengubah sudut perkenaan bola. Scheunemann (2014) menekankan bahwa pemain yang lengkap harus menguasai ketiga teknik soalnya tiap situasi di lapangan butuh jenis tendangan yang berbeda."
)

add_sub_point("d. Rekomendasi untuk Pembelajaran")
add_isi(
    "Dalam pembelajaran PJOK, sebaiknya ketiga teknik diajarkan secara bertahap. Mulai dari kaki bagian dalam yang paling mudah dan paling sering dipakai, kemudian punggung kaki untuk melatih power tendangan, dan terakhir kaki bagian luar sebagai variasi dan pengembangan keterampilan. Yang penting, setiap teknik perlu dilatih dalam konteks permainan sesungguhnya — jangan hanya menendang bola diam, tapi juga menendang bola yang sedang bergerak, menendang sambil berlari, dan menendang di bawah tekanan penjagaan lawan (Luxbacher, 2011)."
)

# ================================================================
# 1.6. RINGKASAN
# ================================================================

add_heading2("1.6.\tRingkasan")

add_isi(
    "Menendang bola adalah keterampilan paling mendasar sekaligus paling penting di sepak bola. Dari pembahasan di bab ini, kita sudah menganalisis tiga teknik tendangan utama yang masing-masing punya karakter dan fungsi tersendiri. Tendangan kaki bagian dalam jadi andalan untuk operan pendek dan menengah karena akurasinya tinggi berkat area perkenaan yang datar dan lebar. Jangan lupa soal posisi kaki tumpuan yang harus tepat di samping bola, pergelangan kaki yang dikunci, dan follow-through yang mengarah ke target."
)
add_isi(
    "Tendangan kaki bagian luar menawarkan variasi yang tidak terduga karena bisa dilakukan tanpa mengubah posisi tubuh secara drastis dan menghasilkan putaran bola ke arah luar. Meski akurasinya lebih menantang, teknik ini sangat berguna dalam situasi-situasi yang membutuhkan kreativitas. Sedangkan tendangan punggung kaki menjadi pilihan utama saat dibutuhkan power besar, misalnya untuk tembakan jarak jauh dan long passing. Kunci dari teknik ini ada di ayunan pinggul yang penuh, pergelangan kaki yang terkunci keras, dan kepala yang tetap menunduk saat menendang. Ketiga teknik ini saling melengkapi dan idealnya dikuasai semua oleh setiap pemain sepak bola."
)

# ================================================================
# 1.7. PENILAIAN
# ================================================================

add_heading2("1.7.\tPenilaian")

add_isi("1) Praktikkan tendangan kaki bagian dalam ke target yang sudah ditentukan sebanyak sepuluh kali, lalu hitung berapa kali bola tepat mengenai target. Analisis faktor apa yang menyebabkan tendangan meleset!")
add_isi("2) Bandingkan hasil tendangan menggunakan kaki bagian dalam dan kaki bagian luar ke target yang sama. Catat perbedaan yang kamu rasakan dari segi kenyamanan, akurasi, dan kekuatan tendangan!")
add_isi("3) Lakukan tendangan punggung kaki ke gawang dari jarak 16 meter, lalu analisis posisi tubuh dan pergelangan kaki — apakah sudah sesuai dengan teknik yang benar atau masih ada kesalahan!")
add_isi("4) Dalam permainan mini sepak bola, coba gunakan ketiga teknik tendangan secara bergantian sesuai situasi. Setelah permainan selesai, diskusikan bersama kelompok tentang teknik mana yang paling sering terpakai dan kenapa!")
add_isi("5) Rekam gerakan menendang bola temanmu dari samping menggunakan handphone, lalu analisis bersama-sama apakah fase persiapan, perkenaan bola, dan follow-through sudah dilakukan dengan benar!")

# ================================================================
# DAFTAR PUSTAKA
# ================================================================
add_page_break()
add_heading1("DAFTAR PUSTAKA")

pustaka = [
    "FIFA. (2023). Laws of the Game 2023/24. Zurich: Fédération Internationale de Football Association.",
    "Luxbacher, J. A. (2011). Sepak Bola: Langkah-Langkah Menuju Sukses (Edisi Keempat). Jakarta: Raja Grafindo Persada.",
    "Mahendra, A. (2020). Pendidikan Jasmani dan Strategi Pembelajaran. Bandung: FPOK UPI Press.",
    "Mielke, D. (2007). Dasar-Dasar Sepak Bola: Cara yang Lebih Baik untuk Mempelajarinya. Bandung: Pakar Raya.",
    "Scheunemann, T. (2014). Kurikulum dan Pedoman Dasar Sepak Bola Indonesia. Jakarta: PSSI.",
    "Sucipto. (2015). Pembelajaran Sepak Bola: Filosofi, Teknik, dan Taktik. Bandung: FPOK UPI Press.",
]

for ref in pustaka:
    p = doc.add_paragraph(style='isi')
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ======== SAVE ========
output_path = r"d:\SUPERVISOR\SPORT\Buku_PJKR_BAB1_Tendangan.docx"
doc.save(output_path)
print(f"Dokumen berhasil disimpan: {output_path}")

# ======== UPDATE TOC VIA WORD COM ========
print("Updating TOC dan nomor halaman via Word...")
import win32com.client, os

word = win32com.client.Dispatch('Word.Application')
word.Visible = False
try:
    abs_path = os.path.abspath(output_path)
    wdoc = word.Documents.Open(abs_path)
    
    # Update all fields
    for story in wdoc.StoryRanges:
        story.Fields.Update()
    
    # Try update TOC
    if wdoc.TablesOfContents.Count > 0:
        wdoc.TablesOfContents(1).Update()
    
    wdoc.Save()
    wdoc.Close()
    print("Berhasil! TOC dan page numbers sudah ter-update.")
except Exception as e:
    print(f"Info: {e}")
    print("Silakan update manual: buka file > Ctrl+A > F9")
finally:
    word.Quit()

print(f"\nFile final: {output_path}")
