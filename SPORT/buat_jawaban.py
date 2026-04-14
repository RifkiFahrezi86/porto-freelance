from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.54)

# ======== STYLES ========
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.paragraph_format.line_spacing = 1.5

# ======== HEADER / IDENTITAS ========
def add_centered(text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p

def add_normal(text, bold=False, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_soal(nomor, teks):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"Soal {nomor}:")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run2 = p2.add_run(teks)
    run2.italic = True
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    return p2

def add_jawaban_header():
    p = doc.add_paragraph()
    run = p.add_run("Jawaban:")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def add_jawaban(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_spacer():
    p = doc.add_paragraph()
    p.space_after = Pt(6)

# ======== COVER ========
add_centered("UJIAN AKHIR SEMESTER", bold=True, size=14)
add_centered("MATA KULIAH: PERKEMBANGAN PESERTA DIDIK", bold=True, size=14)
add_centered("")
add_centered("")
add_centered("Disusun oleh:", size=12)
add_centered("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Nama\t: Nur Arslan")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("NIM\t: 250301502050")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Kelas\t: PJKR K25")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

add_centered("")
add_centered("")
add_centered("PROGRAM STUDI PENDIDIKAN JASMANI, KESEHATAN DAN REKREASI", bold=True, size=12)
add_centered("2026", bold=True, size=12)

doc.add_page_break()

# ======== PETUNJUK ========
add_normal("Petunjuk:", bold=True)
add_normal("Jawablah seluruh pertanyaan secara sistematis, argumentatif, dan berbasis teori. Gunakan contoh kontekstual untuk memperkuat jawaban. Cantumkan rujukan teori jika diperlukan.")
add_spacer()

# ============================================================
# SOAL 1
# ============================================================
add_soal(1, "Jelaskan secara konseptual perbedaan antara pertumbuhan (growth) dan perkembangan (development) dalam konteks peserta didik, serta implikasinya dalam proses pembelajaran!")
add_jawaban_header()

add_jawaban(
    "Pertumbuhan (growth) merujuk pada perubahan kuantitatif yang bisa diukur secara fisik seperti bertambahnya tinggi dan berat badan, sedangkan perkembangan (development) lebih mengarah pada perubahan kualitatif berupa peningkatan fungsi kognitif, emosional, sosial, dan moral individu (Hurlock, 1980). Sebagai contoh, anak yang bertambah tinggi sedang mengalami pertumbuhan, namun ketika ia mulai bisa berpikir logis dan mengendalikan emosi, itu adalah perkembangan. Implikasinya bagi pembelajaran, guru perlu menyadari bahwa kesiapan belajar peserta didik tidak hanya ditentukan oleh usia atau fisiknya saja, tetapi juga oleh kematangan psikologisnya, sehingga pendekatan pembelajaran harus disesuaikan dengan tahap perkembangan masing-masing peserta didik agar proses belajar berjalan efektif dan bermakna."
)
add_spacer()

# ============================================================
# SOAL 2
# ============================================================
add_soal(2, "Uraikan faktor-faktor yang memengaruhi perkembangan peserta didik yang meliputi aspek hereditas dan lingkungan. Analisis bagaimana interaksi kedua faktor tersebut membentuk karakteristik individu!")
add_jawaban_header()

add_jawaban(
    "Perkembangan peserta didik dipengaruhi oleh dua faktor utama yaitu hereditas dan lingkungan. Hereditas mencakup sifat bawaan yang diturunkan lewat gen seperti potensi kecerdasan dan temperamen, sementara lingkungan meliputi pola asuh keluarga, kualitas pendidikan, pergaulan sebaya, dan kondisi sosial-budaya. Menurut Santrock (2019), keduanya tidak bisa dipisahkan karena hereditas menyediakan potensi awal, sedangkan lingkunganlah yang menentukan sejauh mana potensi tersebut berkembang optimal. Misalnya, seorang anak berbakat olahraga tetapi tanpa dukungan fasilitas dan pelatih yang memadai, bakatnya tidak akan terasah dengan baik. Interaksi nature-nurture ini mengajarkan bahwa guru sebaiknya tidak terburu-buru melabeli kemampuan siswa berdasarkan latar belakangnya saja, melainkan berusaha menciptakan lingkungan belajar yang mampu mengoptimalkan beragam potensi setiap anak."
)
add_spacer()

# ============================================================
# SOAL 3
# ============================================================
add_soal(3, "Analisis tahapan perkembangan kognitif menurut Jean Piaget, serta jelaskan relevansinya dalam perancangan strategi pembelajaran di tingkat pendidikan dasar dan menengah!")
add_jawaban_header()

add_jawaban(
    "Piaget membagi perkembangan kognitif anak ke dalam empat tahap: sensorimotor (0-2 tahun) di mana anak memahami dunia lewat indera dan gerakan; praoperasional (2-7 tahun) di mana anak mulai menggunakan simbol namun berpikir masih egosentris; operasional konkret (7-11 tahun) di mana anak sudah mampu berpikir logis terhadap hal-hal nyata seperti konsep konservasi dan klasifikasi; serta operasional formal (11 tahun ke atas) di mana individu sudah bisa berpikir abstrak dan hipotetis. Relevansinya dalam pendidikan, di tingkat SD guru sebaiknya menggunakan media konkret seperti alat peraga dan bermain peran karena siswa masih berada di tahap operasional konkret, sedangkan di tingkat menengah guru bisa mengajak diskusi konsep abstrak dan analisis masalah karena siswa sudah memasuki tahap operasional formal. Intinya, strategi pembelajaran perlu diselaraskan dengan tahap berpikir peserta didik agar materi bisa dipahami secara bermakna."
)
add_spacer()

# ============================================================
# SOAL 4
# ============================================================
add_soal(4, "Jelaskan konsep Zone of Proximal Development (ZPD) menurut Lev Vygotsky dan diskusikan peran scaffolding dalam meningkatkan kemampuan belajar peserta didik!")
add_jawaban_header()

add_jawaban(
    "Zone of Proximal Development (ZPD) menurut Vygotsky adalah jarak antara kemampuan yang sudah dikuasai anak secara mandiri dengan kemampuan yang bisa dicapainya apabila mendapat bantuan dari pihak yang lebih kompeten, baik guru, orang tua, maupun teman sebaya. Dalam kerangka ZPD ini, Bruner memperkenalkan konsep scaffolding, yakni bantuan sementara yang diberikan secara bertahap dan dikurangi seiring meningkatnya kemampuan anak hingga ia akhirnya bisa belajar mandiri. Contoh konkretnya dalam pelajaran olahraga, guru yang mengajarkan servis bola voli awalnya memegangi tangan siswa untuk menunjukkan gerakan yang benar, lalu hanya memberi aba-aba verbal, dan akhirnya membiarkan siswa berlatih sendiri. Pendekatan ini efektif karena peserta didik tetap ditantang untuk berkembang tanpa merasa terbebani oleh tugas yang terlalu sulit."
)
add_spacer()

# ============================================================
# SOAL 5
# ============================================================
add_soal(5, "Bandingkan dan analisis teori perkembangan psikososial Erik Erikson dengan teori perkembangan moral Lawrence Kohlberg dalam memahami dinamika perilaku peserta didik!")
add_jawaban_header()

add_jawaban(
    "Erikson membagi perkembangan psikososial manusia ke dalam delapan tahap krisis, misalnya industry vs. inferiority pada usia 6-12 tahun dan identity vs. role confusion pada remaja 12-18 tahun, sedangkan Kohlberg membagi perkembangan moral ke dalam tiga tingkatan yaitu prakonvensional (orientasi hukuman dan hadiah), konvensional (orientasi persetujuan sosial dan kepatuhan hukum), serta pascakonvensional (orientasi prinsip etika universal). Teori Erikson bersifat lebih luas karena mencakup seluruh aspek psikososial sepanjang hayat, sementara teori Kohlberg lebih spesifik pada penalaran moral. Meskipun berbeda fokus, keduanya saling melengkapi; misalnya remaja yang sedang mencari jati diri menurut Erikson bisa jadi juga sedang beralih dari tingkat konvensional ke pascakonvensional dalam penalaran moralnya menurut Kohlberg. Pemahaman terhadap kedua teori ini membantu guru mendampingi peserta didik bukan hanya secara akademik, tetapi juga dalam membangun karakter dan kematangan moral."
)
add_spacer()

# ============================================================
# SOAL 6
# ============================================================
add_soal(6, "Uraikan karakteristik perkembangan sosial dan emosional pada masa remaja, serta identifikasi faktor-faktor yang dapat memengaruhi kestabilan emosi peserta didik!")
add_jawaban_header()

add_jawaban(
    "Pada masa remaja, secara sosial individu mulai menjadikan kelompok teman sebaya sebagai rujukan utama dalam bersikap, bahkan rela mengubah penampilan demi pengakuan sosial, sementara secara emosional remaja sering mengalami fluktuasi perasaan yang tajam akibat beberapa faktor. Faktor-faktor yang memengaruhi kestabilan emosi antara lain: (1) faktor biologis berupa perubahan hormonal selama pubertas, (2) faktor psikologis terkait pencarian identitas diri, (3) faktor lingkungan keluarga seperti pola asuh yang terlalu otoriter atau permisif, (4) faktor lingkungan sosial termasuk tekanan teman sebaya dan risiko bullying, serta (5) paparan media sosial yang memicu perbandingan sosial dan menurunkan harga diri. David Elkind menyebut fenomena imaginary audience di mana remaja merasa selalu menjadi pusat perhatian semua orang. Guru, terutama guru pendidikan jasmani, perlu menciptakan suasana belajar yang aman secara emosional agar setiap peserta didik merasa dihargai dan tidak takut mengekspresikan dirinya."
)
add_spacer()

# ============================================================
# SOAL 7
# ============================================================
add_soal(7, "Analisis pentingnya perbedaan individu dalam perkembangan peserta didik, serta jelaskan strategi yang dapat dilakukan guru untuk mengakomodasi keberagaman tersebut dalam pembelajaran!")
add_jawaban_header()

add_jawaban(
    "Setiap peserta didik membawa latar belakang, kemampuan, gaya belajar, dan kecepatan perkembangan yang berbeda, dan menurut Gardner melalui teori Multiple Intelligences, kecerdasan manusia bersifat majemuk sehingga tidak bisa diukur dengan satu parameter saja. Pengabaian terhadap perbedaan ini dapat membuat siswa kehilangan motivasi dan mengembangkan sikap negatif terhadap sekolah. Strategi yang bisa dilakukan guru antara lain: menerapkan pembelajaran berdiferensiasi dengan menyesuaikan konten, proses, dan produk pembelajaran; menggunakan metode yang bervariasi seperti ceramah, diskusi, praktik langsung, dan media visual; memberikan pilihan tugas yang memungkinkan siswa menunjukkan kemampuan sesuai kekuatannya; serta melakukan asesmen yang beragam dan tidak hanya mengandalkan tes tertulis. Dalam konteks pendidikan jasmani, guru bisa memberikan penilaian proses yang menghargai usaha dan perkembangan setiap siswa, bukan semata-mata hasil akhir."
)
add_spacer()

# ============================================================
# SOAL 8
# ============================================================
add_soal(8, "Bagaimana pengaruh perkembangan teknologi digital terhadap aspek kognitif, sosial, dan emosional peserta didik dalam konteks pendidikan modern?")
add_jawaban_header()

add_jawaban(
    "Teknologi digital berdampak pada tiga aspek perkembangan peserta didik: secara kognitif, di satu sisi teknologi membuka akses informasi tanpa batas dan mengasah kemampuan berpikir kritis, tetapi di sisi lain memunculkan kebiasaan shallow thinking karena terbiasa mencari jawaban instan; secara sosial, media sosial memperluas jaringan pergaulan namun dapat mengurangi keterampilan komunikasi tatap muka dan meningkatkan risiko cyberbullying; serta secara emosional, penggunaan berlebihan berpotensi memicu kecemasan, gangguan tidur, dan rendahnya harga diri akibat perbandingan sosial yang konstan, meskipun teknologi juga bisa dimanfaatkan positif seperti melalui aplikasi meditasi atau platform konseling daring. Dalam konteks pendidikan modern, peran guru bukan hanya menyampaikan materi tetapi juga menjadi pembimbing yang membantu peserta didik mengembangkan literasi digital dan mengelola penggunaan teknologi secara sehat dan bertanggung jawab."
)
add_spacer()

# ============================================================
# SOAL 9
# ============================================================
add_soal(9, "Jelaskan urgensi pemahaman teori perkembangan peserta didik bagi pendidik dalam merancang pembelajaran yang efektif, adaptif, dan berpusat pada peserta didik!")
add_jawaban_header()

add_jawaban(
    "Pemahaman teori perkembangan sangat penting bagi guru karena memberikan landasan ilmiah untuk merancang pembelajaran yang tepat sasaran. Untuk pembelajaran yang efektif, guru perlu menyesuaikan tingkat kesulitan materi dengan tahap perkembangan kognitif peserta didik agar tidak terlalu mudah atau terlalu sulit. Untuk pembelajaran yang adaptif, pemahaman terhadap perbedaan individu memungkinkan guru menyesuaikan pendekatan dan metode sesuai kebutuhan tanpa memaksakan standar tunggal. Untuk pembelajaran yang berpusat pada peserta didik, teori Erikson mengingatkan bahwa setiap tahap usia membawa tantangan psikososial tertentu, sementara teori Vygotsky menekankan pentingnya interaksi sosial dan dukungan yang tepat. Dengan mengintegrasikan berbagai teori perkembangan dalam praktik mengajar, guru dapat menciptakan lingkungan belajar yang tidak hanya mencerdaskan secara akademik tetapi juga membentuk peserta didik yang sehat secara emosional dan matang secara sosial."
)
add_spacer()

# ============================================================
# SOAL 10
# ============================================================
add_soal(10, "Seorang peserta didik menunjukkan kecenderungan kurang aktif dalam interaksi sosial di kelas, namun memiliki capaian akademik yang baik. Lakukan analisis berdasarkan teori perkembangan yang relevan, dan rumuskan langkah-langkah pedagogis yang dapat dilakukan oleh guru untuk mengoptimalkan perkembangan peserta didik tersebut.")
add_jawaban_header()

add_jawaban(
    "Menurut teori Erikson, peserta didik ini mungkin berhasil mengatasi krisis industry vs. inferiority di ranah akademik tetapi belum mampu membangun rasa percaya dalam hubungan sosialnya, atau jika ia remaja, sedang bergulat dengan krisis identity vs. role confusion. Dari perspektif Vygotsky, kurangnya interaksi sosial bisa menghambat perkembangan kognitif tingkat tinggi karena pengetahuan dibangun lewat interaksi, sedangkan menurut Gardner, anak ini kemungkinan kuat dalam kecerdasan intrapersonal tetapi belum optimal di kecerdasan interpersonal. Langkah-langkah pedagogis yang bisa dilakukan guru antara lain: melakukan observasi dan pendekatan personal untuk memahami akar penyebabnya; menempatkannya dalam kelompok kecil yang suportif dengan peran sebagai tutor sebaya; merancang aktivitas kooperatif dalam pelajaran olahraga yang menuntut kerja sama tim; memberikan penguatan positif setiap kali ia menunjukkan usaha berpartisipasi sosial; serta berkomunikasi dengan orang tua untuk menyusun strategi yang konsisten. Tujuannya bukan mengubah kepribadiannya menjadi ekstrovert, melainkan memastikan ia punya keterampilan sosial yang cukup untuk berfungsi baik dalam berbagai situasi."
)

# ======== DAFTAR PUSTAKA ========
doc.add_page_break()
add_normal("DAFTAR PUSTAKA", bold=True)
add_spacer()

pustaka = [
    "Erikson, E. H. (1968). Identity: Youth and Crisis. New York: W.W. Norton & Company.",
    "Gardner, H. (1983). Frames of Mind: The Theory of Multiple Intelligences. New York: Basic Books.",
    "Hurlock, E. B. (1980). Psikologi Perkembangan: Suatu Pendekatan Sepanjang Rentang Kehidupan (Terjemahan). Jakarta: Erlangga.",
    "Kohlberg, L. (1984). The Psychology of Moral Development. San Francisco: Harper & Row.",
    "Piaget, J. (1972). The Psychology of the Child. New York: Basic Books.",
    "Santrock, J. W. (2019). Life-Span Development (17th ed.). New York: McGraw-Hill Education.",
    "Vygotsky, L. S. (1978). Mind in Society: The Development of Higher Psychological Processes. Cambridge: Harvard University Press.",
]

for ref in pustaka:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ======== SAVE ========
output_path = r"d:\SUPERVISOR\SPORT\Jawaban_UAS_PPD_Nur_Arslan.docx"
doc.save(output_path)
print(f"Dokumen berhasil disimpan di: {output_path}")
